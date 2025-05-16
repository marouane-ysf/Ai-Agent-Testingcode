import streamlit as st
import asyncio
from dotenv import load_dotenv
from azure.identity.aio import DefaultAzureCredential
from azure.ai.projects import AIProjectClient
from semantic_kernel.agents import AzureAIAgent
import re
import PyPDF2
import os
from dotenv import load_dotenv
from fpdf import FPDF
import base64
from datetime import datetime
import io
from docx import Document
from pypdf import PdfReader
import fitz #PyMuPDF
import uuid
from io import BytesIO

# Charger la configuration depuis le fichier .env
load_dotenv()

# Récupération des IDs des agents à partir du fichier .env
AGENT_IDS = {
    "manager": os.environ.get("MANAGER_AGENT_ID"),
    "router": os.environ.get("ROUTER_AGENT_ID"),
    "quality": os.environ.get("QUALITY_AGENT_ID"),
    "drafter": os.environ.get("DRAFT_AGENT_ID"),
    "contracts_compare": os.environ.get("COMPARE_AGENT_ID"),
    "market_comparison": os.environ.get("MarketComparisonAgent_ID"),
    "negotiation": os.environ.get("NegotiationAgent_ID")
}

# Configuration du projet Azure
PROJECT_CONN_STR = os.environ.get("AZURE_AI_PROJECT_CONNECTION_STRING")

# Définition des agents avec leurs informations
AGENTS = {
    "manager": {"name": "Manager Agent", "icon": "🧭", "description": "Répond a des questions d'ordre générale sur le management de contrat et"},
    "quality": {"name": "Agent Qualité", "icon": "🔍", "description": "Évalue la qualité des contrats et identifie les erreurs clés"},
    "drafter": {"name": "Agent Rédacteur", "icon": "📝", "description": "Prépare des ébauches structurées"},
    "contracts_compare": {"name": "Agent Comparaison de Contrats", "icon": "⚖️", "description": "Compare les contrats et fournit un tableau de comparaison"},
    "market_comparison": {"name": "Agent Comparaison de Marché", "icon": "📊", "description": "Compare différentes options de marché et fournit des insights"},
    "negotiation": {"name": "Agent Négociation", "icon": "🤝", "description": "Assiste dans les stratégies et tactiques de négociation"}
}

# Fonction de détection heuristique d'agent basée sur des mots-clés et patterns
def heuristic_agent_selection(query):
    """
    Utilise des heuristiques basées sur des mots-clés pour déterminer l'agent approprié
    quand le router ne fonctionne pas correctement.
    """
    query = query.lower()

    # Mots-clés pour chaque agent
    keywords = {
        "quality": ['analys', 'évalue', 'qualité', 'risque', 'examine', 'erreur', 'faiblesse', 'problème', 'conformité', 'lacune', 'vérifi', 'identifi', 'point faible', 'point fort', 'critique'],
        "drafter": ['rédige', 'rédaction', 'écri', 'ébauche', 'contrat', 'prépar', 'modèle', 'template', 'document', 'structure', 'clause', 'formulaire', 'proposition', 'accord', 'convention'],
        "contracts_compare": ['compar', 'contrat', 'analyse', 'différence', 'similitude', 'contraste', 'évaluation', 'examen', 'point de comparaison', 'élément de comparaison', 
                              'distinction', 'document', 'clause', 'analyse comparative', 'confronter', 'rapprocher', 'mettre en parallèle', 'juxtaposer'],
        "market_comparison": ['compar', 'marché', 'option', 'insight', 'analyse de marché', 'benchmark'],
        "negotiation": ['négoci', 'stratégie', 'tactique', 'accord', 'contrat', 'discussion', 'proposition']
    }

    # Comptage des occurrences de mots-clés
    keyword_counts = {agent: sum(1 for keyword in keywords[agent] if keyword in query) for agent in keywords}

    # Détection de patterns spécifiques
    patterns = {
        "drafter": r'(rédige[rz]?|écri[rstvez]+|prépar[ez]+)\s+([uneod]+\s+)?(ébauche|contrat|document|proposition)',
        "quality": r'(analyse[rz]?|évalue[rz]?|identifi[ez]+|vérifi[ez]+)\s+([uncedo]+\s+)?(contrat|document|qualité|risque)',
        "contracts_compare": r'(compar[eai]+\s+)?(contrat|document|analyse|différence|similitude|contraste)',
        "market_comparison": r'(compar[eai]+\s+)?(marché|option|insight|analyse de marché|benchmark)',
        "negotiation": r'(négoci[eai]+\s+)?(stratégie|tactique|accord|contrat|discussion|proposition)'
    }

    for agent, pattern in patterns.items():
        if re.search(pattern, query):
            return [agent], f"Motif de {AGENTS[agent]['name'].lower()} détecté"

    # Sélection basée sur le comptage des mots-clés
    selected_agents = [agent for agent, count in keyword_counts.items() if count > 0]
    if not selected_agents:
        if any(word in query for word in ['créer', 'faire', 'rédiger', 'écrire', 'préparer']):
            return ["drafter"], "Aucun mot-clé fort, mais intention de création détectée"
        elif any(word in query for word in ['analyser', 'évaluer', 'vérifier', 'examiner']):
            return ["quality"], "Aucun mot-clé fort, mais intention d'analyse détectée"
        else:
            return ["quality"], "Aucun pattern détecté, fallback par défaut"

    return selected_agents, f"Basé sur les occurrences de mots-clés"

# Fonction pour déterminer les agents à utiliser via le Router Agent
async def determine_appropriate_agents(client, query):
    """
    Utilise le Router Agent pour déterminer quels agents spécialisés sont les plus appropriés
    pour répondre à la requête de l'utilisateur.
    """
    heuristic_agents, heuristic_reason = heuristic_agent_selection(query)

    try:
        st.session_state.progress_text = "🧭 Router Agent: Analyse de la requête..."
        st.session_state.progress_value = 0.2

        # Utiliser un thread existant ou en créer un nouveau pour le router
        router_thread_id = st.session_state.get("router_thread_id", None)
        if router_thread_id:
            # Vérifier si le thread existe toujours
            try:
                await client.agents.get_thread(router_thread_id)
            except:
                router_thread_id = None
        
        if not router_thread_id:
            router_thread = await client.agents.create_thread()
            router_thread_id = router_thread.id
            st.session_state.router_thread_id = router_thread_id
        
        router_prompt = f"""
        Analyze this query and determine the most appropriate agents to handle it.
        Respond with a comma-separated list of agent names: quality, drafter, compare, market_comparison, negotiation.

        Query: {query}

        Remember:
        - Choose "quality" for analysis, evaluation, or identification of issues
        - Choose "drafter" for writing, preparing documents, or creating templates
        - Choose "contracts_compare" for information comparison of two or more contracts.
        - Choose "market_comparison" for comparing market options and providing insights
        - Choose "negotiation" for assistance in negotiation strategies and tactics

        Your comma-separated list of agents:
        """

        await client.agents.create_message(
            thread_id=router_thread_id,
            role="user",
            content=router_prompt
        )

        router_run = await client.agents.create_run(
            thread_id=router_thread_id,
            agent_id=AGENT_IDS["router"]
        )

        while True:
            router_run = await client.agents.get_run(thread_id=router_thread_id, run_id=router_run.id)
            if router_run.status == "completed":
                break
            await asyncio.sleep(1)

        messages = await client.agents.list_messages(thread_id=router_thread_id)
        assistant_messages = [m for m in messages.data if m.role == "assistant"]
        selected_agents = []
        raw_response = "Pas de réponse"

        if assistant_messages:
            latest_message = assistant_messages[-1]
            if latest_message.content:
                for content_item in latest_message.content:
                    if content_item.type == "text":
                        raw_response = content_item.text.value.strip().lower()
                        st.session_state.router_raw_response = raw_response

                        selected_agents = [agent.strip() for agent in raw_response.split(",")]

        if not selected_agents or any(agent not in AGENTS for agent in selected_agents):
            st.session_state.progress_text = f"⚠ Router a échoué. Utilisation de l'heuristique: {', '.join(AGENTS[agent]['name'] for agent in heuristic_agents)}"
            st.session_state.progress_value = 0.3
            return heuristic_agents, raw_response, f"Heuristique ({heuristic_reason})"

        st.session_state.progress_text = f"✅ Agents sélectionnés: {', '.join(AGENTS[agent]['name'] for agent in selected_agents)}"
        st.session_state.progress_value = 0.3
        return selected_agents, raw_response, "Router Agent"

    except Exception as e:
        st.error(f"Erreur lors de la détermination des agents: {e}")
        return heuristic_agents, f"Erreur: {str(e)}", f"Heuristique ({heuristic_reason})"

# Fonction pour exécuter un agent
async def execute_agent(client, agent_id, agent_info, message_content):
    """
    Exécute un agent spécifique avec un message donné et retourne sa réponse.
    """
    try:
        st.session_state.progress_text = f"{agent_info['icon']} {agent_info['name']}: Traitement en cours..."

        # Créer un nouveau thread à chaque fois pour éviter les problèmes
        thread = await client.agents.create_thread()
        thread_id = thread.id
        
        # Ajouter l'historique des conversations si le mode contexte est activé
        if st.session_state.get("context_mode", True) and "messages" in st.session_state:
            # Ajouter l'historique des derniers messages au thread
            context_messages = st.session_state.messages[-3:]  # Prendre les 3 derniers messages
            history = "\n\nHistorique de conversation:\n"
            
            for msg in context_messages:
                if msg["role"] == "user":
                    history += f"Question: {msg['content']}\n"
                else:
                    # Simplifier la réponse pour éviter que l'agent ne se répète
                    resp = msg["content"]
                    if len(resp) > 200:
                        resp = resp[:200] + "..."
                    history += f"Réponse précédente: {resp}\n"
            
            # Ajouter l'historique au message actuel
            if context_messages:
                enhanced_message = f"{history}\n\nNouvelle question: {message_content}"
            else:
                enhanced_message = message_content
        else:
            enhanced_message = message_content

        # Créer le message avec le contexte
        await client.agents.create_message(
            thread_id=thread_id,
            role="user",
            content=enhanced_message
        )

        run = await client.agents.create_run(
            thread_id=thread_id,
            agent_id=agent_id
        )

        while True:
            run = await client.agents.get_run(thread_id=thread_id, run_id=run.id)
            if run.status == "completed":
                break
            await asyncio.sleep(1)

        messages = await client.agents.list_messages(thread_id=thread_id)
        assistant_messages = [m for m in messages.data if m.role == "assistant"]
        response = "Pas de réponse de l'agent"

        if assistant_messages:
            latest_message = assistant_messages[-1]
            response = ""
            if latest_message.content:
                for content_item in latest_message.content:
                    if content_item.type == "text":
                        response += content_item.text.value

        return response

    except Exception as e:
        error_msg = f"Erreur lors de l'exécution de {agent_info['name']}: {e}"
        st.error(error_msg)
        return error_msg
        
# Fonction pour obtenir une analyse de qualité
async def execute_quality_analysis(client, selected_agents, user_query):
    """
    Si l'agent sélectionné n'est pas Quality, exécute une analyse rapide
    avec l'Agent Quality pour améliorer le contexte pour les autres agents.
    """
    if "quality" in selected_agents:
        return None

    try:
        st.session_state.progress_text = "🔍 Analyse préliminaire avec Agent Qualité..."
        st.session_state.progress_value = 0.4

        analysis = await execute_agent(
            client,
            AGENT_IDS["quality"],
            AGENTS["quality"],
            f"Faites une analyse rapide et concise de cette requête: {user_query}"
        )

        return analysis

    except Exception as e:
        st.error(f"Erreur lors de l'analyse préliminaire: {e}")
        return None

# Fonction pour exécuter le workflow complet
async def run_orchestrated_workflow(query):
    """
    Exécute le workflow complet avec orchestration intelligente
    """
    try:
        credential = DefaultAzureCredential()
        project_client = AIProjectClient.from_connection_string(
            conn_str=PROJECT_CONN_STR,
            credential=credential
        )

        async with credential, AzureAIAgent.create_client(credential=credential) as client:
            selected_agents, router_response, selection_method = await determine_appropriate_agents(client, query)
            st.session_state.selected_agents = selected_agents

            analysis = await execute_quality_analysis(client, selected_agents, query)

            responses = {}
            combined_response = ""

            for agent in selected_agents:
                st.session_state.progress_text = f"{AGENTS[agent]['icon']} {AGENTS[agent]['name']}: Traitement en cours..."
                st.session_state.progress_value = 0.6

                prompt = query
                if analysis:
                    prompt = f"En tenant compte de cette analyse: {analysis}\n\n{query}"

                response = await execute_agent(
                    client,
                    AGENT_IDS[agent],
                    AGENTS[agent],
                    prompt
                )

                responses[agent] = response
                combined_response += f"{AGENTS[agent]['name']}:\n{response}\n\n"

            st.session_state.progress_text = "✅ Traitement terminé"
            st.session_state.progress_value = 1.0

            return {
                "selected_agents": selected_agents,
                "agent_names": [AGENTS[agent]['name'] for agent in selected_agents],
                "agent_icons": [AGENTS[agent]['icon'] for agent in selected_agents],
                "combined": combined_response,
                "router_response": router_response,
                "selection_method": selection_method,
                **responses
            }

    except Exception as e:
        return {"error": f"Erreur lors de l'exécution du workflow: {str(e)}"}

# Fonction pour exécuter un pipeline séquentiel
async def run_sequential_pipeline(query):
    """
    Exécute un pipeline séquentiel avec les agents définis par l'utilisateur
    """
    try:
        credential = DefaultAzureCredential()
        project_client = AIProjectClient.from_connection_string(
            conn_str=PROJECT_CONN_STR,
            credential=credential
        )

        async with credential, AzureAIAgent.create_client(credential=credential) as client:
            sequence = st.session_state.get("agent_sequence", [])
            if not sequence:
                return {"error": "Aucune séquence d'agents définie. Veuillez définir une séquence dans la barre latérale."}

            responses = {}
            current_input = query

            for i, agent_key in enumerate(sequence):
                agent_info = AGENTS[agent_key]
                st.session_state.progress_text = f"{agent_info['icon']} {agent_info['name']}: Traitement en cours..."
                st.session_state.progress_value = (i + 1) / len(sequence)

                response = await execute_agent(
                    client,
                    AGENT_IDS[agent_key],
                    agent_info,
                    current_input
                )

                responses[agent_key] = response
                current_input = f"Tenant compte de la réponse précédente: {response}\n\nQuestion initiale: {query}"  # Ajoute le contexte pour l'agent suivant

            st.session_state.progress_text = "✅ Traitement terminé"
            st.session_state.progress_value = 1.0

            combined_response = "\n\n".join(f"Réponse de {AGENTS[agent_key]['name']}:\n{response}" for agent_key, response in responses.items())

            return {
                "selected_agent": "sequence",
                "agent_name": "Multi-Agent Séquentiel",
                "agent_icon": "🔄",
                "combined": combined_response,
                **responses
            }

    except Exception as e:
        return {"error": f"Erreur lors de l'exécution du workflow multi-agent: {str(e)}"}

# Fonction pour exécuter un agent spécifique
async def run_specific_agent(query, agent_key):
    """
    Exécute un agent spécifique (mode agent unique)
    """
    try:
        credential = DefaultAzureCredential()
        project_client = AIProjectClient.from_connection_string(
            conn_str=PROJECT_CONN_STR,
            credential=credential
        )

        async with credential, AzureAIAgent.create_client(credential=credential) as client:
            agent_id = AGENT_IDS[agent_key]
            agent_name = AGENTS[agent_key]["name"]
            agent_icon = AGENTS[agent_key]["icon"]

            st.session_state.progress_text = f"{agent_icon} {agent_name}: Préparation de votre réponse..."
            st.session_state.progress_value = 0.5

            response = await execute_agent(client, agent_id, AGENTS[agent_key], query)

            st.session_state.progress_text = "✅ Traitement terminé"
            st.session_state.progress_value = 1.0

            return {
                "selected_agent": agent_key,
                "agent_name": agent_name,
                "agent_icon": agent_icon,
                "combined": response
            }

    except Exception as e:
        return {"error": f"Erreur lors de l'exécution de l'agent {agent_key}: {str(e)}"}

# Fonction pour exécuter les fonctions asynchrones dans Streamlit
def run_async_function(func, *args, **kwargs):
    loop = asyncio.new_event_loop()
    asyncio.set_event_loop(loop)
    try:
        return loop.run_until_complete(func(*args, **kwargs))
    finally:
        loop.close()

# Fonction pour extraire le texte d'un PDF avec OCR
def extract_text_from_pdf_ocr(pdf_document):
    """for OCR"""
    text = ""
    for page_num in range(pdf_document.page_count):
        page = pdf_document.load_page(page_num)
        text += page.get_text()
    return text

# Fonction pour extraire le texte d'un fichier PDF ou texte
def extract_text_from_pdf(uploaded_file, ocr):
    if (uploaded_file is not None) and ocr:
        file_name = uploaded_file.name
        pdf_document = fitz.open(stream=uploaded_file.read(), filetype="pdf")
        extracted_text = extract_text_from_pdf_ocr(pdf_document)
        raw_text = extracted_text
        return raw_text, file_name
    else:
        if uploaded_file is not None:
            file_name = uploaded_file.name
            if uploaded_file.type == "text/plain":
                raw_text = str(uploaded_file.read(),"utf-8")
            elif uploaded_file.type == "application/pdf":
                reader = PdfReader(uploaded_file)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n"
                raw_text = text

        return raw_text, file_name

def generate_pdf_with_fitz(title, content):
    """
    Génère un document PDF à partir du contenu fourni en utilisant PyMuPDF (fitz)
    pour une meilleure prise en charge des caractères Unicode
    """
    # Créer un nouveau document PDF
    doc = fitz.open()
    page = doc.new_page()
    
    # Ajouter le titre
    font_size = 16
    text_rect = fitz.Rect(50, 50, page.rect.width - 50, 70)
    page.insert_text(text_rect.tl, title, fontsize=font_size, fontname="helvetica-bold")
    
    # Ajouter le contenu
    y_position = 100
    paragraphs = content.split('\n')
    for paragraph in paragraphs:
        if paragraph.strip():
            if paragraph.strip().startswith('#'):
                # Titre
                text = paragraph.strip('# ')
                text_rect = fitz.Rect(50, y_position, page.rect.width - 50, y_position + 20)
                page.insert_text(text_rect.tl, text, fontsize=14, fontname="helvetica-bold")
                y_position += 25
            else:
                # Texte normal
                text_rect = fitz.Rect(50, y_position, page.rect.width - 50, y_position + 20)
                # Découper les lignes trop longues
                words = paragraph.split()
                line = ""
                for word in words:
                    test_line = line + " " + word if line else word
                    if len(test_line) > 80:  # Limite de caractères par ligne
                        page.insert_text(fitz.Point(50, y_position), line, fontsize=12, fontname="helvetica")
                        y_position += 15
                        line = word
                    else:
                        line = test_line
                
                if line:  # Insérer la dernière ligne
                    page.insert_text(fitz.Point(50, y_position), line, fontsize=12, fontname="helvetica")
                    y_position += 20
        else:
            y_position += 10
            
    # Ajouter le pied de page avec la date
    footer_text = f"Document généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')}"
    footer_rect = fitz.Rect(50, page.rect.height - 30, page.rect.width - 50, page.rect.height - 10)
    page.insert_text(footer_rect.tl, footer_text, fontsize=10, fontname="helvetica-oblique")
    
    # Enregistrer dans BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer.getvalue()

def generate_pdf(title, content):
    """
    Génère un document PDF à partir du contenu fourni en utilisant FPDF
    """
    try:
        # Essayer d'abord avec PyMuPDF qui gère mieux l'Unicode
        return generate_pdf_with_fitz(title, content)
    except Exception as e:
        # Fallback sur FPDF en cas d'erreur
        try:
            pdf = FPDF()
            pdf.add_page()
            
            # Configuration du PDF
            pdf.set_font("Arial", "B", 16)
            
            # Encoder le titre pour éviter les problèmes avec les caractères spéciaux
            pdf.cell(0, 10, title.encode('latin-1', 'replace').decode('latin-1'), ln=True, align="C")
            pdf.ln(10)
            
            # Ajout du contenu principal
            pdf.set_font("Arial", "", 12)
            
            # Diviser le contenu en paragraphes
            paragraphs = content.split('\n')
            for paragraph in paragraphs:
                if paragraph.strip():
                    # Vérifier si c'est un titre (commence par #)
                    if paragraph.strip().startswith('#'):
                        pdf.set_font("Arial", "B", 14)
                        cleaned_text = paragraph.strip('# ').encode('latin-1', 'replace').decode('latin-1')
                        pdf.cell(0, 10, cleaned_text, ln=True)
                        pdf.set_font("Arial", "", 12)
                    else:
                        # Texte normal
                        cleaned_text = paragraph.encode('latin-1', 'replace').decode('latin-1')
                        pdf.multi_cell(0, 6, cleaned_text)
                        pdf.ln(4)
            
            # Date de génération
            pdf.set_font("Arial", "I", 10)
            date_text = f"Document généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')}".encode('latin-1', 'replace').decode('latin-1')
            pdf.cell(0, 10, date_text, ln=True, align="R")
            
            # Retourner le PDF comme bytes
            return pdf.output(dest="S").encode('latin-1')
        except Exception as inner_e:
            # En cas d'échec des deux méthodes, créer un PDF très simple
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", "", 12)
            pdf.cell(0, 10, "Error creating formatted PDF. Please see text version.", ln=True)
            return pdf.output(dest="S").encode('latin-1')

def generate_docx(title, content):
    """
    Génère un document DOCX à partir du contenu fourni
    """
    doc = Document()
    
    # Ajouter le titre
    doc.add_heading(title, 0)
    
    # Ajouter le contenu
    paragraphs = content.split('\n')
    for paragraph in paragraphs:
        if paragraph.strip():
            # Vérifier si c'est un titre (commence par #)
            if paragraph.strip().startswith('#'):
                level = paragraph.count('#')
                doc.add_heading(paragraph.strip('# '), level)
            else:
                # Texte normal
                doc.add_paragraph(paragraph)
    
    # Date de génération
    doc.add_paragraph(f"Document généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')}", style='Subtitle')
    
    # Retourner le document comme bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def create_download_button(content, filename, button_text):
    """
    Crée un bouton de téléchargement personnalisé pour le contenu fourni
    """
    # Vérifier si le contenu est en bytes, sinon le convertir
    if not isinstance(content, bytes):
        content = content.encode()
        
    b64 = base64.b64encode(content).decode()
    button_uuid = str(uuid.uuid4()).replace('-', '')
    button_id = f'download-button-{button_uuid}'
    
    custom_css = f"""
        <style>
            #{button_id} {{
                background-color: #4CAF50;
                color: white;
                padding: 10px 15px;
                border: none;
                border-radius: 5px;
                cursor: pointer;
                font-size: 14px;
                text-align: center;
                text-decoration: none;
                display: inline-block;
                margin: 4px 2px;
            }}
            #{button_id}:hover {{
                background-color: #45a049;
            }}
        </style>
    """
    
    dl_link = custom_css + f'<a download="{filename}" id="{button_id}" href="data:application/octet-stream;base64,{b64}">{button_text}</a><br></br>'
    
    return dl_link

# Fonction pour extraire le texte de plusieurs fichiers
def extract_text_from_multiple_files(uploaded_files, ocr):
    """
    extract text from multiple files and return a list of file text
    """
    files_text = []

    # Create a progress bar
    progress_text = "extracting files content..."
    progress_bar = st.progress(0)

    total_files = len(uploaded_files)
    for i, uploaded_file in enumerate(uploaded_files):
        # Update progress
        progress_bar.progress(i / total_files)

        # Extract the content
        file_content, file_name = extract_text_from_pdf(uploaded_file, ocr)
        if file_name:
            files_text.append({
                'content': file_content,
                'name': file_name
            })
            st.write(f"✅ {uploaded_file.name} content extracted successfully")
        else:
            st.error(f"❌ Failed to extract content from {uploaded_file.name}")

    # Complete the progress bar
    progress_bar.progress(1.0)

    return files_text

def prompt_constructor(user_input, ocr):
    msg, files = user_input["text"], user_input["files"]
    if files:
        if msg is None:
            msg = "sharing documents"
        files_content = extract_text_from_multiple_files(files, ocr)
        user_prompt = msg
        for i, file in enumerate(files_content):
            user_prompt += f"\ncontract n°{i+1} called " + file["name"] + "\n" + file["content"]
            # Vérifiez que st.session_state.uploaded_file est une liste
            if "uploaded_file" in st.session_state and isinstance(st.session_state.uploaded_file, list):
                st.session_state.uploaded_file.append(file)
            else:
                st.session_state.uploaded_file = [file]
    else:
        user_prompt = msg
        
    # Ajout d'un résumé de l'historique pour maintenir le contexte
    if "messages" in st.session_state and len(st.session_state.messages) > 0:
        context = "Voici l'historique de notre conversation :\n"
        for msg in st.session_state.messages[-5:]:  # Utiliser les 5 derniers messages
            role = "Utilisateur" if msg["role"] == "user" else "Assistant"
            content = msg["content"]
            context += f"{role}: {content}\n"
        
        user_prompt = f"{context}\n\nNouvelle demande: {user_prompt}"
    
    return user_prompt