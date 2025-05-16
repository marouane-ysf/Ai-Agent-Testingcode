"""
Utilitaires pour la génération de fichiers (PDF, DOCX, etc.).
"""

import io
import fitz  # PyMuPDF
from docx import Document
from datetime import datetime
from fpdf import FPDF
from typing import Union

def generate_pdf_with_fitz(title: str, content: str) -> bytes:
    """
    Génère un document PDF à partir du contenu fourni en utilisant PyMuPDF (fitz).
    
    Args:
        title: Titre du document
        content: Contenu du document
        
    Returns:
        bytes: Contenu PDF en bytes
    """
    # Créer un nouveau document PDF
    doc = fitz.open()
    page = doc.new_page()
    
    # Ajouter le titre
    font_size = 16
    text_rect = fitz.Rect(50, 50, page.rect.width - 50, 70)
    page.insert_text(text_rect.tl, title, fontsize=font_size, fontname="helvetica-bold")
    
    # Ajouter le contenu
    y_position = 100
    paragraphs = content.split('\n')
    for paragraph in paragraphs:
        if paragraph.strip():
            if paragraph.strip().startswith('#'):
                # Titre
                text = paragraph.strip('# ')
                text_rect = fitz.Rect(50, y_position, page.rect.width - 50, y_position + 20)
                page.insert_text(text_rect.tl, text, fontsize=14, fontname="helvetica-bold")
                y_position += 25
            else:
                # Texte normal
                text_rect = fitz.Rect(50, y_position, page.rect.width - 50, y_position + 20)
                # Découper les lignes trop longues
                words = paragraph.split()
                line = ""
                for word in words:
                    test_line = line + " " + word if line else word
                    if len(test_line) > 80:  # Limite de caractères par ligne
                        page.insert_text(fitz.Point(50, y_position), line, fontsize=12, fontname="helvetica")
                        y_position += 15
                        line = word
                    else:
                        line = test_line
                
                if line:  # Insérer la dernière ligne
                    page.insert_text(fitz.Point(50, y_position), line, fontsize=12, fontname="helvetica")
                    y_position += 20
        else:
            y_position += 10
            
    # Ajouter le pied de page avec la date
    footer_text = f"Document généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')}"
    footer_rect = fitz.Rect(50, page.rect.height - 30, page.rect.width - 50, page.rect.height - 10)
    page.insert_text(footer_rect.tl, footer_text, fontsize=10, fontname="helvetica-oblique")
    
    # Enregistrer dans BytesIO
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    
    return buffer.getvalue()

def generate_pdf(title: str, content: str) -> bytes:
    """
    Génère un document PDF à partir du contenu fourni (avec fallback sur FPDF).
    
    Args:
        title: Titre du document
        content: Contenu du document
        
    Returns:
        bytes: Contenu PDF en bytes
    """
    try:
        # Essayer d'abord avec PyMuPDF qui gère mieux l'Unicode
        return generate_pdf_with_fitz(title, content)
    except Exception as e:
        # Fallback sur FPDF
        try:
            pdf = FPDF()
            pdf.add_page()
            
            # Configuration du PDF
            pdf.set_font("Arial", "B", 16)
            
            # Encoder le titre pour éviter les problèmes avec les caractères spéciaux
            pdf.cell(0, 10, title.encode('latin-1', 'replace').decode('latin-1'), ln=True, align="C")
            pdf.ln(10)
            
            # Ajout du contenu principal
            pdf.set_font("Arial", "", 12)
            
            # Diviser le contenu en paragraphes
            paragraphs = content.split('\n')
            for paragraph in paragraphs:
                if paragraph.strip():
                    # Vérifier si c'est un titre (commence par #)
                    if paragraph.strip().startswith('#'):
                        pdf.set_font("Arial", "B", 14)
                        cleaned_text = paragraph.strip('# ').encode('latin-1', 'replace').decode('latin-1')
                        pdf.cell(0, 10, cleaned_text, ln=True)
                        pdf.set_font("Arial", "", 12)
                    else:
                        # Texte normal
                        cleaned_text = paragraph.encode('latin-1', 'replace').decode('latin-1')
                        pdf.multi_cell(0, 6, cleaned_text)
                        pdf.ln(4)
            
            # Date de génération
            pdf.set_font("Arial", "I", 10)
            date_text = f"Document généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')}".encode('latin-1', 'replace').decode('latin-1')
            pdf.cell(0, 10, date_text, ln=True, align="R")
            
            # Retourner le PDF comme bytes
            return pdf.output(dest="S").encode('latin-1')
        except Exception as inner_e:
            # En cas d'échec des deux méthodes, créer un PDF très simple
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", "", 12)
            pdf.cell(0, 10, "Error creating formatted PDF. Please see text version.", ln=True)
            return pdf.output(dest="S").encode('latin-1')

def generate_docx(title: str, content: str) -> bytes:
    """
    Génère un document DOCX à partir du contenu fourni.
    
    Args:
        title: Titre du document
        content: Contenu du document
        
    Returns:
        bytes: Contenu DOCX en bytes
    """
    doc = Document()
    
    # Ajouter le titre
    doc.add_heading(title, 0)
    
    # Ajouter le contenu
    paragraphs = content.split('\n')
    for paragraph in paragraphs:
        if paragraph.strip():
            # Vérifier si c'est un titre (commence par #)
            if paragraph.strip().startswith('#'):
                level = paragraph.count('#')
                doc.add_heading(paragraph.strip('# '), level)
            else:
                # Texte normal
                doc.add_paragraph(paragraph)
    
    # Date de génération
    doc.add_paragraph(f"Document généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')}", style='Subtitle')
    
    # Retourner le document comme bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()